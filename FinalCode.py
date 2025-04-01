from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from base64 import urlsafe_b64decode, urlsafe_b64encode
from datetime import datetime, timedelta
import re
import time
import pandas as pd
from tabulate import tabulate
import logging
import io
import os
from docx import Document  # For processing Word documents
from PyPDF2 import PdfReader  # For processing PDFs

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
def get_emails_by_job_id(service, job_id):
    """Fetch all emails containing a given job ID within the last 30 days."""
    today = datetime.now()
    last_month = today - timedelta(days=30)
    last_month_str = last_month.strftime('%Y/%m/%d')

    all_messages = []
    page_token = None

    while True:
        try:
            results = service.users().messages().list(
                userId="me",
                q=f"after:{last_month_str} {job_id}",
                maxResults=500,
                pageToken=page_token
            ).execute()

            messages = results.get("messages", [])
            all_messages.extend(messages)

            page_token = results.get("nextPageToken")
            if not page_token:
                break

            time.sleep(1)  # Avoid rate limiting

        except Exception as e:
            logging.error(f"Error fetching emails: {e}")
            break

    return all_messages
def decode_base64(data):
    """Decodes base64 email content safely."""
    missing_padding = len(data) % 4
    if missing_padding:
        data += '=' * (4 - missing_padding)
    return urlsafe_b64decode(data)
def extract_email_body(payload):
    """Extracts the email body content."""
    if not payload:
        return ""

    if "body" in payload and "data" in payload["body"]:
        return decode_base64(payload["body"]["data"]).decode("utf-8", errors="ignore")

    if "parts" in payload:
        for part in payload["parts"]:
            if part.get("mimeType", "") in ["text/plain", "text/html"] and "data" in part.get("body", {}):
                return decode_base64(part["body"]["data"]).decode("utf-8", errors="ignore")
            if "parts" in part:
                nested_body = extract_email_body(part)
                if nested_body:
                    return nested_body

    return ""
def extract_skill_names(skills):
    """Extracts skill names from skill sentences and preserves skill type."""
    skill_names = []
    for skill in skills:
        # Extract skill type (Required, Highly desired, Nice to have)
        skill_type_match = re.search(
            r"(Required|Highly desired|Nice to have)\s+\d+\s+Months",
            skill,
            re.IGNORECASE
        )
        skill_type = skill_type_match.group(0) if skill_type_match else ""

        # Extract skill name using regex (customize as needed)
        skill_name_match = re.search(
            r"(?:Ability\s+to\s+|Understanding\s+|Strong\s+knowledge\s+of\s+|Experience\s+or\s+interest\s+in\s+|Ability\s+to\s+learn\s+|Excellent\s+|Knowledge\s+of\s+)(.*?)\s*(?:Required|Highly desired|Nice to have|\d+\s+Months|$)",
            skill,
            re.IGNORECASE
        )
        if skill_name_match:
            skill_name = skill_name_match.group(1).strip()
            # Capitalize the first letter of the skill name
            skill_name = skill_name[0].upper() + skill_name[1:]
            skill_names.append(f"• {skill_name} {skill_type}")
        else:
            skill_names.append(f"• {skill}")  # Fallback to the full sentence if no skill name is found

    return skill_names
def extract_skills(email_body):
    """Extracts Skills section while preserving 'Required X Years' text."""
    match = re.search(
        r"Skills?:\s*(.*?)\s*(?:Responsibilities:|Qualifications:|Description:|Job ID:|$)",
        email_body,
        re.DOTALL | re.IGNORECASE
    )
    if not match:
        return None

    skill_text = match.group(1).strip()
    return clean_skill_text(skill_text)
def clean_skill_text(skill_text):
    """Cleans formatting while retaining skill text with years."""
    if not skill_text:
        return None

    # Remove numbering/bullets while preserving other text
    skill_text = re.sub(r"(\n\s*[-•*]\s*\d+\.?\s*|\n\s*[-•*]\s*)", "\n", skill_text)
    skill_text = re.sub(r"^\d+\.\s*", "", skill_text, flags=re.MULTILINE)

    # Split into individual skills based on newlines or specific patterns
    skills = []
    current_skill = ""
    for line in skill_text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Detect new skill (keywords like "Ability", "Understanding", "Strong", etc.)
        if re.match(r"^(Ability|Understanding|Strong|Experience|Knowledge|Excellent)", line, re.IGNORECASE):
            if current_skill:
                skills.append(current_skill.strip())
            current_skill = line
        else:
            current_skill += " " + line

    if current_skill:
        skills.append(current_skill.strip())

    # Extract skill names and preserve skill type
    skill_names = extract_skill_names(skills)

    return skill_names
def extract_details_from_body(body):
    """Extracts comprehensive candidate details from email body."""
    # Extract name
    name_patterns = [
        r"First Name\s*\(.*?\):\s*(.*?)\s*Middle Name\s*\(.*?\):\s*(.*?)\s*Last Name\s*\(.*?\):\s*(.*?)(?:\n|$)",
        r"Name\s*:\s*(.*?)(?:\n|$)",
        r"<(b|strong)>(.*?)</\1>"
    ]
 
    name = "N/A"
    for pattern in name_patterns:
        match = re.search(pattern, body, re.IGNORECASE | re.DOTALL)
        if match:
            # Extract groups and clean up
            groups = [g.strip() for g in match.groups() if g]
            # Remove ">" characters and clean up the name
            cleaned_groups = []
            for g in groups:
                # Replace ">" and any spaces around it with a single space
                cleaned = re.sub(r'\s*>\s*', ' ', g)
                cleaned_groups.append(cleaned)
            name = " ".join(filter(None, cleaned_groups))
            break
 
    # Extract contact information
    phone_match = re.search(
        r"(?i)(?:phone\s*#|phone\s*no|phone|ph|mobile|contact)[\s#:]*([+]?[\d]{0,4}[\s\-.]*\(?\d{0,4}\)?[\s\-.]*\d{3,4}[\s\-.]*\d{3,4})",
        body
    )
    phone = phone_match.group(1).strip() if phone_match else "N/A"
 
    email_match = re.search(
        r"(?i)email\s*[:\-]?\s*([\w\.-]+@[\w\.-]+\.\w+)",
        body
    )
    email = email_match.group(1).strip() if email_match else "N/A"
 
    location_match = re.search(
        r"(?i)Current location\s*(?:\(city/state\))?\s*[:\-]?\s*(?:[\*\b_]{1,2})?\s*([\w\s.,-]+(?:[/,]\s*[\w\s.-]+)?)\s*(?:[\*\b_]{1,2})?\s*(?=\n|$)",
        body
    )
    current_location = location_match.group(1).strip() if location_match else "N/A"
 
    # Extract experience
    experience = "0 years"
    experience_patterns = [
        r"Total no of years experience:\s*([\d]+(?:\.\d+)?\+?)",  # Matches "13+", "5+", "10", "7.5"
        r"Years of experience:\s*([\d]+(?:\.\d+)?\+?)",  # Common pattern
        r"Experience:\s*([\d]+(?:\.\d+)?\+?)",  # Generic pattern
        r"Experience:\s*([\d.+]*)"
    ]
    for pattern in experience_patterns:
        match = re.search(pattern, body, re.IGNORECASE)
        if match:
            exp_value = match.group(1).strip()
            experience = f"{exp_value} years" if not exp_value.lower().endswith("years") else exp_value
            break
 
    # Extract professional details
    certification_count = 0
    cert_match = re.search(r"Certification Count:\s*(\d+)", body, re.IGNORECASE)
    if cert_match:
        certification_count = int(cert_match.group(1).strip())
 
    govt_exp_match = re.search(
    r"(?i)Government\s*experience\s*:\s*(?:\(mention\s*the\s*government\s*name's\s*in\s*resume\s*otherwise\s*No\)\s*:?\s*)?([\w\s,&.\-]+?(?:&\w+;)?[\w\s,&.\-]*?)(?:\n|$|\r\n)",
    body,
    re.IGNORECASE
    )
    # Extract as a string, not a list
    if govt_exp_match:
        government_experience = govt_exp_match.group(1).strip()
        government_experience = re.sub(r'\s+', ' ', government_experience)
        if government_experience.lower() in ["no", "not mentioned", "none"]:
           government_experience = "Not worked with the government"
    else:
       government_experience = "Not worked with the government"

    # Extract visa info
    visa_info = "N/A"
    visa_patterns = [
        r"Visa\s*Status\s*with\s*Validity\s*:\s*([^\n\r]*)",  # Matches "Visa Status with Validity: GC EAD, 2029"
        r"Visa\s*type\s*and\s*sponsor\s*name\s*\(.*?\)\s*:\s*([^\n\r]*)",  # Matches "Visa type and sponsor name (ex: H1 - ABC Inc): H1B"
        r"Visa\s*type\s*:\s*([^\n\r]*)",  # Matches "Visa type: H1B"
        r"Status\s*:\s*([^\n\r]*)"  # Matches "Status: GC EAD"
    ]
 
    visa_info = "N/A"  # Default value
    for pattern in visa_patterns:
        match = re.search(pattern, body, re.IGNORECASE)
        if match:
            visa_info = match.group(1).strip()
            if not visa_info or visa_info.isspace():  # If there's no visa status name after colon, assign "N/A"
                visa_info = "N/A"
            break
 
    return {
        "Candidate Name": name,
        "Phone No": phone,
        "Email": email,
        "Current Location": current_location,
        "Total Experience": experience,
        "Certification Count": certification_count,
        "Government Experience": government_experience,
        "Visa Status": visa_info.strip()
    }
def extract_text_from_attachment(attachment_data, filename):
    try:
        if filename.endswith(".pdf"):
            # Extract text from PDF
            reader = PdfReader(io.BytesIO(attachment_data))
            text = "\n".join(page.extract_text() for page in reader.pages)
            return text
        elif filename.endswith(".docx"):
            # Extract text from Word document
            doc = Document(io.BytesIO(attachment_data))
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        else:
            logging.warning(f"Unsupported file type: {filename}")
            return None
    except Exception as e:
        logging.error(f"Error extracting text from {filename}: {e}")
        return None
def process_attachments(service, message_id):
    """Collects valid (filename, attachment_id) tuples."""
    try:
        msg = service.users().messages().get(userId="me", id=message_id, format="full").execute()
        payload = msg.get("payload", {})
        parts = payload.get("parts", [])

        attachments = []
        for part in parts:
            filename = part.get("filename")
            attachment_id = part.get("body", {}).get("attachmentId")
            # Only include entries with BOTH filename and attachment ID
            if filename and attachment_id:
                attachments.append((filename, attachment_id))

        return attachments

    except Exception as e:
        logging.error(f"Error listing attachments: {e}")
        return []
def get_attachment_data(service, message_id, attachment_id):
    """Fetches and decodes attachment data using the Gmail API."""
    try:
        attachment = service.users().messages().attachments().get(
            userId="me",
            messageId=message_id,
            id=attachment_id
        ).execute()
        data = attachment.get("data", "")
        return urlsafe_b64decode(data)
    except Exception as e:
        logging.error(f"Error fetching attachment data: {e}")
        return None
def filter_excluded_files(attachments):
    exclusion_terms = ["rtr", "sow", "sm", "jd", "job", "description", "mail", "contract", "project", "h1", "gc", "dl", "signed"]
    filtered = []
    for filename in attachments:
        lower_name = filename.lower()
        # Skip files with exclusion terms
        if any(term in lower_name for term in exclusion_terms):
            continue
        # Keep only PDF/DOCX files
        if filename.lower().endswith((".pdf", ".docx")):
            filtered.append(filename)
    return filtered
def is_resume_content(text):
    required_sections = [
        r"work\s*experience", r"education", r"skills", r"summary", r"projects", r"certifications",
        r"expertise\s*brief", r"professional\s*engagements", r"academic\s*qualifications",
        r"certificates\s*&\s*accolades", r"key\s*skills", r"executive\s*summary", r"technical\s*skills",
        r"professional\s*summary", r"career\s*highlights", r"professional\s*milestones",
        r"responsibilities", r"qualification\s*badges", r"technical\s*profile",
        r"significant\s*practices", r"certification", r"technical\s*expertise", r"client",
        r"education\s*&\s*credentials", r"business\s*process\s*improvement", r"career\s*achievements",
        r"summary\s*of\s*the\s*experience", r"work\s*/\s*assignment\s*history",
        r"education\s*and\s*professional\s*qualifications", r"work\s*/\s*assignment\s*history"
    ]
    found = sum(1 for section in required_sections if re.search(section, text, re.IGNORECASE))
    return found >= 1  # Adjusted threshold to 1 for broader detection
def identify_resume(service, message_id, attachments):
    """Identifies the resume from pre-validated attachments."""
    # Step 1: Filter out excluded files and non-PDF/DOCX
    valid_files = filter_excluded_files([fn for (fn, _) in attachments])

    # Step 2: Check content of remaining files
    for filename in valid_files:
        # Get attachment ID from pre-validated list
        attachment_id = next((aid for (fn, aid) in attachments if fn == filename), None)
        if not attachment_id:
            continue

        attachment_data = get_attachment_data(service, message_id, attachment_id)
        if not attachment_data:
            continue

        # Extract text and validate
        text = extract_text_from_attachment(attachment_data, filename)
        if text and is_resume_content(text):
            return filename

    return "N/A"
import re

def extract_skills_from_subject(subject):
    """
    Extracts skills dynamically from the email subject, ignoring any prefix.
    Returns the skills as a Python list.
    """
    # Define a generic regex pattern to match the skills part after any prefix
    # Assumes the skills part starts after a keyword like "with", "including", or similar
    skill_pattern = r"(?:.*\bwith\b\s*|\bincluding\b\s*)(.*)"

    # Search for the skills part in the subject
    match = re.search(skill_pattern, subject, re.IGNORECASE)
    
    if match:
        # Extract the skills part
        skills_part = match.group(1).strip()
        
        # Split the skills part into individual skills (assuming they are separated by commas)
        skills_list = [skill.strip() for skill in skills_part.split(",")]
        
        # Return the skills as a list
        return skills_list
    else:
        # If no match, return the entire subject as a single-item list (fallback)
        return [subject]
def extract_email_data(service, message_id):
    """Fetches the email and extracts details, skills, and subject."""
    try:
        msg = service.users().messages().get(userId="me", id=message_id, format="full").execute()

        # Extract the subject from the email headers
        headers = msg.get("payload", {}).get("headers", [])
        subject = next((header["value"] for header in headers if header["name"] == "Subject"), "N/A")

        # Extract skills from the email subject dynamically
        subject_skills = extract_skills_from_subject(subject)  # Use the updated function

        email_body = extract_email_body(msg.get("payload", {}))

        if not email_body:
            return None, None, None

        # Extract skills from email body
        email_skills = extract_skills(email_body)
        details = extract_details_from_body(email_body)

        # Process attachments if any
        attachment_results = process_attachments(service, message_id)
        filenames = [fn for (fn, _) in attachment_results]

        # Identify resume (ignore excluded files, validate content)
        resume_filename = identify_resume(service, message_id, attachment_results)
        details["Resume File"] = resume_filename

        return details, email_skills, subject_skills  # Return the extracted skills as a list

    except Exception as e:
        logging.error(f"Error processing email {message_id}: {e}")
        return None, None, None
import os
import shutil

def create_resume_folder(folder_name="Resumes"):
    """
    Creates a folder to store resumes. If the folder already exists, clears all files in it.
    
    Args:
        folder_name (str): Name of the folder to create.
    
    Returns:
        str: Path to the folder.
    """
    if os.path.exists(folder_name):
        # Clear all files in the folder
        for filename in os.listdir(folder_name):
            file_path = os.path.join(folder_name, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)  # Delete the file
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)  # Delete the subdirectory
            except Exception as e:
                logging.error(f"Failed to delete {file_path}: {e}")
        logging.info(f"Cleared existing files in folder: {folder_name}")
    else:
        # Create the folder if it doesn't exist
        os.makedirs(folder_name)
        logging.info(f"Created folder: {folder_name}")
    
    return folder_name
def save_resumes_to_folder(service, details, message_id, attachments, resume_folder):
    """
    Saves resumes to the specified folder.
    
    Args:
        service: Gmail API service object.
        details (dict): Dictionary containing candidate details, including the resume filename.
        message_id: The message ID of the email.
        attachments: List of attachments from the email.
        resume_folder (str): Folder to save the resume files.
    """
    resume_filename = details["Resume File"]
    if resume_filename == "N/A":
        # logging.info(f"No resume file found for candidate: {details['Candidate Name']}")
        return

    # Find the attachment ID for the resume file
    attachment_id = next((aid for (fn, aid) in attachments if fn == resume_filename), None)
    if not attachment_id:
        # logging.warning(f"Attachment ID not found for resume file: {resume_filename}")
        return

    # Fetch attachment data
    attachment_data = get_attachment_data(service, message_id, attachment_id)
    if not attachment_data:
        # logging.error(f"Failed to fetch attachment data for resume file: {resume_filename}")
        return

    # Save the resume file to the folder
    resume_file_path = os.path.join(resume_folder, resume_filename)
    try:
        with open(resume_file_path, "wb") as file:
            file.write(attachment_data)
        # logging.info(f"Saved resume file: {resume_file_path}")
    except Exception as e:
        logging.error(f"Failed to save resume file {resume_filename}: {e}")
def extract_text_from_resume(file_path):
    """
    Extracts text from a resume file (PDF or DOCX).
    """
    try:
        if file_path.endswith(".pdf"):
            # Extract text from PDF
            reader = PdfReader(file_path)
            text = "\n".join(page.extract_text() for page in reader.pages)
            return text
        elif file_path.endswith(".docx"):
            # Extract text from Word document
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        else:
            logging.warning(f"Unsupported file type: {file_path}")
            return None
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {e}")
        return None
def extract_skills_from_resume_text(text):
    if not text:
        return []

    # Step 1: Define common "Skills" section headers
    skills_section_headers = [
        r"Technical Skills",
        r"Skills",
        r"Core Competencies",
        r"Key Skills",
        r"Technical Expertise",
        r"Professional Skills",
    ]

    # Step 2: Try to locate the "Skills" section using the headers
    skills_section_text = ""
    for header in skills_section_headers:
        match = re.search(
            rf"(?i){header}[:\\s]*(.*?)(?:\\n\\n|\\n[A-Z]|$)",
            text,
            re.DOTALL,
        )
        if match:
            skills_section_text = match.group(1).strip()
            break

    # Step 3: If no "Skills" section is found, use the entire text
    if not skills_section_text:
        skills_section_text = text

    # Debugging: Print the skills section text
    print(f"Skills Section Text from Resume:\n{skills_section_text[:500]}")  # Print first 500 characters

    # Step 4: Extract skills using regex patterns
    skill_patterns = [
        r"(?i)\\b(?:python|java|c\\+\\+|javascript|sql|html|css|react|angular|node\\.?js|aws|docker|kubernetes|machine learning|data science|tensorflow|pytorch|nlp|artificial intelligence|ai|agile|scrum|devops|git|jenkins|ansible|terraform)\\b",
        r"(?i)\\b(?:project management|cloud computing|cybersecurity|big data|spark|hadoop|tableau|power bi|nosql|mongodb|postgresql|mysql|linux|unix|rest api|graphql|microservices|azure|gcp|selenium|jira)\\b",
    ]

    validated_skills = set()
    for pattern in skill_patterns:
        skills = re.findall(pattern, skills_section_text, re.IGNORECASE)
        validated_skills.update(skills)

    # Debugging: Print extracted skills
    print(f"Extracted Skills from Resume: {list(validated_skills)}")

    return list(validated_skills)
def update_resume_skills(df, resume_folder):
    """
    Updates the DataFrame with skills extracted from the resumes in the specified folder.
    """
    resume_skills_list = []

    for index, row in df.iterrows():
        resume_filename = row["Resume File"]
        if resume_filename == "N/A":
            resume_skills_list.append("N/A")
            continue

        # Construct the full path to the resume file
        resume_file_path = os.path.join(resume_folder, resume_filename)

        # Extract text from the resume file
        try:
            with open(resume_file_path, "rb") as file:
                attachment_data = file.read()
            resume_text = extract_text_from_attachment(attachment_data, resume_filename)
            if not resume_text:
                logging.warning(f"Failed to extract text from resume file: {resume_filename}")
                resume_skills_list.append("N/A")
                continue

            # Extract skills from the resume text
            skills = extract_skills_from_resume_text(resume_text)
            resume_skills_list.append(", ".join(skills) if skills else "N/A")
        except Exception as e:
            logging.error(f"Error processing resume file {resume_filename}: {e}")
            resume_skills_list.append("N/A")

    # Add the "Resume Skills" column to the DataFrame
    df["Resume Skills"] = resume_skills_list
    return df
from docx import Document
from PyPDF2 import PdfReader
import re
import os

def extract(path):
    """
    Extracts text from a DOCX or PDF file.
    """
    if path.endswith(".docx"):
        # Load the DOCX document
        doc = Document(path)
        # Concatenate all paragraphs into the variable
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        return full_text
    
    elif path.endswith(".pdf"):
        # Load the PDF document
        reader = PdfReader(path)
        # Extract text from all pages
        full_text = ""
        for page in reader.pages:
            full_text += page.extract_text() + "\n"
        return full_text
    
    else:
        raise ValueError("Unsupported file type. Only DOCX and PDF files are supported.")

def match(full_text, subject_skills):
    """
    Matches skills from the subject with the resume text and counts their occurrences.
    Returns a dictionary of {skill: count}.
    """
    text_lower = full_text.lower()
    skill_counts = {}
 
    for skill in subject_skills:
        # Escape special characters and match whole words
        pattern = r'\b' + re.escape(skill.strip().lower()) + r'\b'
        matches = re.findall(pattern, text_lower)
        count = len(matches)
        if count > 0:
            skill_counts[skill.strip()] = count  # Preserve original casing
 
    return skill_counts


def compare_skills_with_resumes(subject_skills, resume_folder):
    """
    Compares subject skills with resume content and returns a dictionary of 
    {resume_filename: {skill: count}}.
    """
    matched_skills_dict = {}
 
    for filename in os.listdir(resume_folder):
        if filename.endswith((".docx", ".pdf")):
            file_path = os.path.join(resume_folder, filename)
            try:
                full_text = extract(file_path)
                skill_counts = match(full_text, subject_skills)
                if skill_counts:
                    matched_skills_dict[filename] = skill_counts
            except Exception as e:
                print(f"Error processing {filename}: {e}")
 
    return matched_skills_dict


def calculate_resume_score(matched_skills, government_experience):
    score = 0
    
    # Calculate points for skills
    if matched_skills and isinstance(matched_skills, str):
        matched_skills_list = [skill.strip() for skill in matched_skills.split(",")]
        score += len(matched_skills_list) * 5  # 5 points per matched skill
    
    # Calculate points for government experience
    if government_experience.lower() != "not worked with the government":
        # Count the number of government experiences (separated by commas)
        government_count = government_experience.count(",") + 1
        score += government_count * 10  # 10 points per government experience
    
    return score



def main():
    """Main function to authenticate and extract job-related details."""
    creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    service = build("gmail", "v1", credentials=creds)
 
    job_id = input("Enter the job ID to search for: ").strip()
    messages = get_emails_by_job_id(service, job_id)
 
    if not messages:
        print(f"No emails found related to job ID {job_id}.")
        return
 
    print(f"Found {len(messages)} emails related to job ID {job_id}:")
 
    extracted_skills_list = []
    email_data = []
    subject_skills_list = []
 
    # Create a folder to store resumes
    resume_folder = create_resume_folder("Resumes")
 
    for msg in messages:
        message_id = msg["id"]
        details, skills, subject_skills = extract_email_data(service, message_id)
 
        if skills:
            extracted_skills_list.extend(skills)
            subject_skills_list.append({"Subject": subject_skills, "Skills": skills})
        elif details:
            email_data.append(details)
            # Process attachments for the current email
            attachments = process_attachments(service, message_id)
            if attachments:
                # Save resumes to the folder
                save_resumes_to_folder(service, details, message_id, attachments, resume_folder)
 
    # Display extracted skills and subject separately (above the table)
    if subject_skills_list:
        print("\n=== Extracted Skills from Subjects (From Emails That Mention Skills) ===")
        for item in subject_skills_list:
            print(f"Subject: {item['Subject']}")
            print("\n=== Extracted Skills (From Emails That Mention Skills) ===")
            for skill in item["Skills"]:
                print(f"- {skill}")
            print("-" * 50)  # Separator for readability
 
    # Compare skills with resumes and calculate scores
    if subject_skills_list and email_data:
        df = pd.DataFrame(email_data)
        subject_skills = subject_skills_list[0]["Subject"]  # Use first email's skills
        matched_skills_dict = compare_skills_with_resumes(subject_skills, resume_folder)
 
        # Build matched skills list with counts
        matched_skills_list = []
        for index, row in df.iterrows():
            resume_filename = row["Resume File"]
            if resume_filename in matched_skills_dict:
                skills_str = ", ".join([
                    f"{skill} ({count})"
                    for skill, count in matched_skills_dict[resume_filename].items()
                ])
                matched_skills_list.append(skills_str)
            else:
                matched_skills_list.append("N/A")
 
        # Add columns to DataFrame
        df["Matched Skills"] = matched_skills_list
       
        # Cleanup columns
        if "Resume Skills" in df.columns:
            df.drop(columns=["Resume Skills"], inplace=True)
 
        # Calculate Resume Score for each candidate
        resume_scores = []
        for index, row in df.iterrows():
            matched_skills = row["Matched Skills"]
            government_experience = row["Government Experience"]
            score = calculate_resume_score(matched_skills, government_experience)
            resume_scores.append(score)
 
        # Add scoring columns
        df["Resume Score"] = resume_scores
        df["Rank"] = df["Resume Score"].rank(ascending=False, method="min").astype(int)
       
        # Sort and display
        df = df.sort_values(by="Rank")
        columns_order = ["Rank"] + [col for col in df.columns if col != "Rank"]
        df = df[columns_order]
 
        print("\n=== Final Candidate Matches with Scores ===")
        print(tabulate(df, headers="keys", tablefmt="grid", showindex=False))
    else:
        print("\nNo skills or resumes found for comparison.")
 
if __name__ == "__main__":
    main()
2025-03-24 21:54:01,363 - INFO - file_cache is only supported with oauth2client<4.0.0
2025-03-24 21:54:41,462 - INFO - Cleared existing files in folder: Resumes
Found 2 emails related to job ID GA-759848:

=== Extracted Skills from Subjects (From Emails That Mention Skills) ===
Subject: ['Articulate Storyline 360', 'digital training', 'TechSmith SnagIt and Adobe experience']

=== Extracted Skills (From Emails That Mention Skills) ===
- • Demonstratable experience using Articulate Storyline 360 and Articulate Studio. Required 3 Years Demonstratable experience developing digital training courses and materials. Required 3 Years
- • Experience editing videos, graphics, sound editors and working with effective instructional design practices Required 3 Years
--------------------------------------------------

=== Final Candidate Matches with Scores ===
+--------+------------------+--------------+----------------------+--------------------+--------------------+-----------------------+--------------------------------+---------------+-------------------------+------------------+----------------+
|   Rank | Candidate Name   | Phone No     | Email                | Current Location   | Total Experience   |   Certification Count | Government Experience          | Visa Status   | Resume File             | Matched Skills   |   Resume Score |
+========+==================+==============+======================+====================+====================+=======================+================================+===============+=========================+==================+================+
|      1 | Wendy Burrowes   | 678-826-7348 | wsburrowes@gmail.com | Atlanta, GA        | 15+ years          |                     0 | Not worked with the government | Citizenship   | WendyB Instr Design.pdf | N/A              |              5 |
+--------+------------------+--------------+----------------------+--------------------+--------------------+-----------------------+--------------------------------+---------------+-------------------------+------------------+----------------+

 